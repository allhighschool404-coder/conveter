from flask import Flask, render_template, request, jsonify, send_file
import pandas as pd
from docx import Document
import io
import re
import fitz
import zipfile
import base64
import os
import tempfile
from PIL import Image
from concurrent.futures import ThreadPoolExecutor
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Table, TableStyle, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.lib import colors

app = Flask(__name__)

SUPPORTED_EXTENSIONS = [
    '.jpeg', '.jpg', '.png', '.webp', '.svg', '.gif', '.psd',
    '.ai', '.eps', '.pdf', '.doc', '.docx', '.txt', '.rtf',
    '.odt', '.xls', '.xlsx', '.csv', '.xlsm', '.ods'
]

# Max worker threads for parallel file & page rendering
MAX_WORKERS = min(16, (os.cpu_count() or 4) * 2)

# Pre-compiled regex patterns for maximum performance
ENGLISH_DATE_PATTERN = re.compile(r'\b\d{1,4}[/\-\.]\d{1,2}[/\-\.]\d{1,4}\b')
BANGLA_DATE_PATTERN = re.compile(r'\b[০-৯]{1,4}[/\-\.][০-৯]{1,2}[/\-\.][০-৯]{1,4}\b')
SUBJECT_SPLIT_PATTERN = re.compile(r'বিষয়|বিষয়|Subject', flags=re.IGNORECASE)
SUB_CLEAN_SPLIT_PATTERN = re.compile(r'মাস|শ্রেণি|শ্রেণী|শিক্ষক|\n')
CLEAN_FILENAME_PATTERN = re.compile(r'[\\/*?:"<>|]')
WHITESPACE_PATTERN = re.compile(r'\s+')

def clean_and_normalize(text):
    if not text:
        return ""
    return " ".join(str(text).strip().lower().split())

def remove_dates_from_text(text):
    """টেক্সটের ভেতর থেকে যেকোনো ফরম্যাটের তারিখ মুছে ফেলার দ্রুত ফাংশন"""
    if not text:
        return ""
    text = str(text).strip()
    text = ENGLISH_DATE_PATTERN.sub('', text)
    text = BANGLA_DATE_PATTERN.sub('', text)
    return text.strip()

def extract_subject_name(doc, original_filename=""):
    """ওয়ার্ড ফাইল থেকে নিখুঁতভাবে বিষয়ের নাম বের করা"""
    subject_name = ""
    try:
        # ১. টেবিলের ভেতর বিষয় খোঁজা
        for table in doc.tables:
            for row in table.rows:
                row_str = " ".join([cell.text.strip() for cell in row.cells if cell.text])
                if any(k in row_str for k in ['বিষয়', 'বিষয়', 'Subject']):
                    parts = SUBJECT_SPLIT_PATTERN.split(row_str)
                    if len(parts) > 1:
                        target = parts[1].replace(':', '').replace('ঃ', '').strip()
                        sub_clean = SUB_CLEAN_SPLIT_PATTERN.split(target)[0].strip()
                        if sub_clean:
                            subject_name = sub_clean
                            break
            if subject_name:
                break

        # ২. সাধারণ প্যারাগ্রাফে বিষয় খোঁজা
        if not subject_name:
            for para in doc.paragraphs:
                text = para.text.strip()
                if any(k in text for k in ['বিষয়', 'বিষয়', 'Subject']):
                    parts = SUBJECT_SPLIT_PATTERN.split(text)
                    if len(parts) > 1:
                        target = parts[1].replace(':', '').replace('ঃ', '').strip()
                        sub_clean = SUB_CLEAN_SPLIT_PATTERN.split(target)[0].strip()
                        if sub_clean:
                            subject_name = sub_clean
                            break
    except Exception as e:
        print(f"Subject extraction error: {e}")

    if not subject_name and original_filename:
        subject_name = original_filename.rsplit('.', 1)[0]

    subject_name = WHITESPACE_PATTERN.sub(' ', subject_name)
    subject_name = CLEAN_FILENAME_PATTERN.sub('', subject_name).strip()
    return subject_name if subject_name else "Lesson_Plan"

def extract_data_strictly_all_pages(docx_file, template_columns, original_filename=""):
    doc = Document(docx_file)
    data = []
    
    subject_title = extract_subject_name(doc, original_filename)
    
    for table in doc.tables:
        header_row_index = -1
        table_matrix = []
        
        for row in table.rows:
            row_text = [cell.text.strip() for cell in row.cells]
            table_matrix.append(row_text)
            
        for index, row in enumerate(table_matrix):
            row_str = "".join([clean_and_normalize(c) for c in row])
            if 'তারিখ' in row_str and 'বার' in row_str:
                header_row_index = index
                break
                
        if header_row_index != -1:
            word_headers = table_matrix[header_row_index]
            
            chapter_idx, syllabus_idx, page_idx, test_idx = -1, -1, -1, -1
            
            for idx, h in enumerate(word_headers):
                h_clean = clean_and_normalize(h)
                if 'অধ্যায়' in h_clean or 'টপিক' in h_clean: chapter_idx = idx
                elif 'পাঠ্যাংশ' in h_clean or 'আলোচ্য' in h_clean: syllabus_idx = idx
                elif 'পৃষ্ঠা' in h_clean: page_idx = idx
                elif 'টেস্ট' in h_clean: test_idx = idx

            for row_index in range(header_row_index + 1, len(table_matrix)):
                current_row = table_matrix[row_index]
                
                if not any(current_row):
                    continue
                
                row_str_check = "".join([clean_and_normalize(c) for c in current_row])
                if any(key in row_str_check for key in ['বিদ্যালয়', 'বিদ্যালয়', 'পরিকল্পনা', 'शिक्षকের নাম']):
                    continue
                    
                row_dict = {col: "" for col in template_columns}
                
                word_chapter = remove_dates_from_text(current_row[chapter_idx]) if (chapter_idx != -1 and chapter_idx < len(current_row)) else ""
                word_syllabus = remove_dates_from_text(current_row[syllabus_idx]) if (syllabus_idx != -1 and syllabus_idx < len(current_row)) else ""
                word_page = remove_dates_from_text(current_row[page_idx]) if (page_idx != -1 and page_idx < len(current_row)) else ""
                word_test = remove_dates_from_text(current_row[test_idx]) if (test_idx != -1 and test_idx < len(current_row)) else ""
                
                for t_col in template_columns:
                    t_clean = clean_and_normalize(t_col)
                    
                    if 'date' in t_clean or 'day' in t_clean:
                        row_dict[t_col] = "" 
                    elif 'chapter' in t_clean or 'unit' in t_clean:
                        row_dict[t_col] = word_chapter
                    elif 'syllabus' in t_clean or 'topics' in t_clean:
                        row_dict[t_col] = word_syllabus
                    elif 'page' in t_clean or 'reference' in t_clean:
                        row_dict[t_col] = word_page
                    elif 'is class test' in t_clean:
                        row_dict[t_col] = word_test
                    elif 'optional' in t_clean or 'title' in t_clean:
                        if word_test.strip():
                            row_dict[t_col] = word_chapter
                        else:
                            row_dict[t_col] = ""

                data.append(row_dict)
                
    return data, subject_title

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/process', methods=['POST'])
def process_files():
    if 'word_file' not in request.files or 'excel_template' not in request.files:
        return jsonify({'error': 'দয়া করে দুটি ফাইলই আপলোড করুন'}), 400
        
    word_file = request.files['word_file']
    excel_template = request.files['excel_template']
    
    try:
        template_df = pd.read_excel(excel_template)
        columns = template_df.columns.tolist()
        
        processed_data, subject_title = extract_data_strictly_all_pages(word_file, columns, word_file.filename)
        
        if not processed_data:
            return jsonify({'error': 'ওয়ার্ড ফাইলের টেবিল থেকে কোনো তথ্য পাওয়া যায়নি।'}), 400
            
        return jsonify({
            'columns': columns,
            'data': processed_data,
            'subject_title': subject_title
        })
        
    except Exception as e:
        print(f"Error processing files: {e}")
        return jsonify({'error': f"প্রসেসিং ত্রুটি: {str(e)}"}), 500

@app.route('/download', methods=['POST'])
def download_file():
    try:
        content = request.json
        columns = content.get('columns', [])
        data = content.get('data', [])
        subject_title = content.get('subject_title', 'Lesson_Plan')
        
        df = pd.DataFrame(data, columns=columns)
        
        output = io.BytesIO()
        with pd.ExcelWriter(output, engine='openpyxl') as writer:
            df.to_excel(writer, index=False)
        output.seek(0)
        
        return send_file(
            output,
            mimetype='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
            as_attachment=True,
            download_name=f"{subject_title}.xlsx"
        )
    except Exception as e:
        return jsonify({'error': f"ডাউনলোড ত্রুটি: {str(e)}"}), 500

# --- FAST PARALLEL CONVERSION ENGINES ---

def process_fitz_doc(doc, dpi=150, quality=85):
    """PyMuPDF ডকুমেন্ট প্যারালাল থ্রেডিং এর মাধ্যমে হাই-স্পিডে জেপিজিতে রেন্ডার করার ফাংশন"""
    total_pages = len(doc)
    pages_data = [None] * total_pages

    def render_page(i):
        page = doc.load_page(i)
        pix = page.get_pixmap(dpi=dpi, alpha=False)
        img_bytes = pix.tobytes("jpeg", jpg_quality=quality)
        img_filename = f"page_{i + 1}.jpg"
        b64_str = base64.b64encode(img_bytes).decode('utf-8')
        pdata = {
            'page': i + 1,
            'filename': img_filename,
            'data_url': f"data:image/jpeg;base64,{b64_str}",
            'width': pix.width,
            'height': pix.height
        }
        return i, img_filename, img_bytes, pdata

    if total_pages > 1:
        with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
            results = list(executor.map(render_page, range(total_pages)))
    else:
        results = [render_page(0)]

    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_STORED) as zip_file:
        for idx, img_filename, img_bytes, pdata in sorted(results, key=lambda x: x[0]):
            zip_file.writestr(img_filename, img_bytes)
            pages_data[idx] = pdata

    doc.close()
    zip_buffer.seek(0)
    zip_b64 = base64.b64encode(zip_buffer.getvalue()).decode('utf-8')
    return pages_data, zip_b64

def process_pillow_image(file_bytes, dpi=150, quality=85):
    """Pillow দিয়ে মাল্টি-ফ্রেমে বা সিঙ্গেল ইমেজ হাই-স্পিডে প্রসেস করা"""
    img = Image.open(io.BytesIO(file_bytes))
    try:
        n_frames = getattr(img, 'n_frames', 1)
    except Exception:
        n_frames = 1

    def process_frame(i):
        t_img = Image.open(io.BytesIO(file_bytes))
        try:
            t_img.seek(i)
        except EOFError:
            t_img.close()
            return None
            
        frame = t_img.copy()
        t_img.close()

        if frame.mode in ('RGBA', 'LA', 'P'):
            frame = frame.convert('RGBA')
            background = Image.new('RGB', frame.size, (255, 255, 255))
            background.paste(frame, mask=frame.split()[3])
            rgb_img = background
        else:
            rgb_img = frame.convert('RGB')

        out_buf = io.BytesIO()
        rgb_img.save(out_buf, format='JPEG', quality=quality, dpi=(dpi, dpi))
        img_bytes = out_buf.getvalue()
        
        img_filename = f"page_{i + 1}.jpg" if n_frames > 1 else "converted_image.jpg"
        b64_str = base64.b64encode(img_bytes).decode('utf-8')
        
        return i, img_filename, img_bytes, {
            'page': i + 1,
            'filename': img_filename,
            'data_url': f"data:image/jpeg;base64,{b64_str}",
            'width': rgb_img.width,
            'height': rgb_img.height
        }

    if n_frames > 1:
        with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
            results = [res for res in executor.map(process_frame, range(n_frames)) if res is not None]
    else:
        results = [process_frame(0)]

    results.sort(key=lambda x: x[0])
    pages_data = []
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_STORED) as zip_file:
        for _, img_filename, img_bytes, pdata in results:
            zip_file.writestr(img_filename, img_bytes)
            pages_data.append(pdata)

    img.close()
    zip_buffer.seek(0)
    zip_b64 = base64.b64encode(zip_buffer.getvalue()).decode('utf-8')
    return pages_data, zip_b64

def process_pillow_to_pdf(file_bytes):
    img = Image.open(io.BytesIO(file_bytes))
    try:
        n_frames = getattr(img, 'n_frames', 1)
    except Exception:
        n_frames = 1

    pdf_doc = fitz.open()
    for i in range(n_frames):
        try:
            img.seek(i)
        except EOFError:
            break
            
        frame = img.copy()
        if frame.mode in ('RGBA', 'LA', 'P'):
            frame = frame.convert('RGBA')
            background = Image.new('RGB', frame.size, (255, 255, 255))
            background.paste(frame, mask=frame.split()[3])
            rgb_img = background
        else:
            rgb_img = frame.convert('RGB')

        out_buf = io.BytesIO()
        rgb_img.save(out_buf, format='JPEG', quality=85)
        img_data = out_buf.getvalue()
        
        page = pdf_doc.new_page(width=rgb_img.width, height=rgb_img.height)
        page.insert_image(fitz.Rect(0, 0, rgb_img.width, rgb_img.height), stream=img_data)

    img.close()
    pdf_bytes = pdf_doc.tobytes(clean=True, deflate=True)
    pdf_doc.close()
    return pdf_bytes

def convert_doc_via_word(file_bytes, ext):
    import pythoncom, win32com.client
    pythoncom.CoInitialize()
    word = None
    pdf_path = None
    src_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as f:
            f.write(file_bytes)
            src_path = f.name
        pdf_path = src_path + ".pdf"
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = False
        doc = word.Documents.Open(src_path)
        doc.SaveAs(pdf_path, FileFormat=17)  # 17 = wdFormatPDF
        doc.Close(False)
        with open(pdf_path, "rb") as pf:
            return pf.read()
    finally:
        if word:
            try: word.Quit()
            except: pass
        pythoncom.CoUninitialize()
        if src_path and os.path.exists(src_path): os.remove(src_path)
        if pdf_path and os.path.exists(pdf_path): os.remove(pdf_path)

def convert_sheet_via_excel(file_bytes, ext):
    import pythoncom, win32com.client
    pythoncom.CoInitialize()
    excel = None
    pdf_path = None
    src_path = None
    try:
        with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as f:
            f.write(file_bytes)
            src_path = f.name
        pdf_path = src_path + ".pdf"
        excel = win32com.client.Dispatch("Excel.Application")
        excel.Visible = False
        excel.DisplayAlerts = False
        wb = excel.Workbooks.Open(src_path)
        wb.ExportAsFixedFormat(0, pdf_path)  # 0 = xlTypePDF
        wb.Close(False)
        with open(pdf_path, "rb") as pf:
            return pf.read()
    finally:
        if excel:
            try: excel.Quit()
            except: pass
        pythoncom.CoUninitialize()
        if src_path and os.path.exists(src_path): os.remove(src_path)
        if pdf_path and os.path.exists(pdf_path): os.remove(pdf_path)

def fallback_txt_to_pdf(txt_bytes):
    try:
        doc = fitz.open(stream=txt_bytes, filetype="txt")
        pdf_bytes = doc.convert_to_pdf()
        doc.close()
        return pdf_bytes
    except Exception:
        pdf_buf = io.BytesIO()
        doc = SimpleDocTemplate(pdf_buf, pagesize=letter, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
        styles = getSampleStyleSheet()
        text = txt_bytes.decode('utf-8', errors='ignore')
        paragraphs = []
        for line in text.split('\n'):
            if line.strip():
                paragraphs.append(Paragraph(line.strip(), styles['Normal']))
            else:
                paragraphs.append(Spacer(1, 10))
        doc.build(paragraphs)
        pdf_buf.seek(0)
        return pdf_buf.getvalue()

def fallback_doc_to_pdf(docx_bytes, ext):
    if ext == '.docx':
        try:
            doc_obj = Document(io.BytesIO(docx_bytes))
            pdf_buf = io.BytesIO()
            doc = SimpleDocTemplate(pdf_buf, pagesize=letter, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
            styles = getSampleStyleSheet()
            elements = []
            for p in doc_obj.paragraphs:
                if p.text.strip():
                    elements.append(Paragraph(p.text.strip(), styles['Normal']))
            for table in doc_obj.tables:
                table_data = []
                for row in table.rows:
                    table_data.append([cell.text.strip() for cell in row.cells])
                if table_data:
                    t = Table(table_data)
                    t.setStyle(TableStyle([
                        ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#4F46E5')),
                        ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
                        ('GRID', (0,0), (-1,-1), 0.5, colors.grey)
                    ]))
                    elements.append(t)
            if elements:
                doc.build(elements)
                pdf_buf.seek(0)
                return pdf_buf.getvalue()
        except Exception as e:
            print(f"Fallback docx error: {e}")
    return None

def fallback_sheet_to_pdf(sheet_bytes, ext):
    try:
        if ext == '.csv':
            df = pd.read_csv(io.BytesIO(sheet_bytes))
        else:
            df = pd.read_excel(io.BytesIO(sheet_bytes))
            
        pdf_buf = io.BytesIO()
        doc = SimpleDocTemplate(pdf_buf, pagesize=letter, leftMargin=36, rightMargin=36, topMargin=36, bottomMargin=36)
        styles = getSampleStyleSheet()
        
        headers = [str(c) for c in df.columns]
        rows = [[str(val) if pd.notna(val) else '' for val in row] for row in df.values]
        table_data = [headers] + rows
        
        t = Table(table_data)
        t.setStyle(TableStyle([
            ('BACKGROUND', (0,0), (-1,0), colors.HexColor('#4F46E5')),
            ('TEXTCOLOR', (0,0), (-1,0), colors.whitesmoke),
            ('GRID', (0,0), (-1,-1), 0.5, colors.grey),
            ('FONTNAME', (0,0), (-1,0), 'Helvetica-Bold'),
            ('FONTSIZE', (0,0), (-1,-1), 9),
        ]))
        
        doc.build([t])
        pdf_buf.seek(0)
        return pdf_buf.getvalue()
    except Exception as e:
        print(f"Fallback sheet error: {e}")
        return None

# --- ENGINE 1: CONVERT TO JPG ---

def convert_any_file_to_jpg(file_bytes, filename, dpi=150, quality=85):
    _, ext = os.path.splitext(filename)
    ext = ext.lower()
    if not ext:
        ext = '.pdf'

    if ext in ['.pdf', '.svg']:
        fitz_type = 'pdf' if ext == '.pdf' else 'svg'
        doc = fitz.open(stream=file_bytes, filetype=fitz_type)
        return process_fitz_doc(doc, dpi, quality)

    if ext in ['.jpeg', '.jpg', '.png', '.webp', '.gif', '.psd']:
        try:
            return process_pillow_image(file_bytes, dpi, quality)
        except Exception as e:
            try:
                doc = fitz.open(stream=file_bytes)
                return process_fitz_doc(doc, dpi, quality)
            except Exception:
                raise RuntimeError(f"ইমেজ ফাইল প্রসেস করতে ব্যর্থ হয়েছে: {str(e)}")

    if ext in ['.ai', '.eps']:
        try:
            doc = fitz.open(stream=file_bytes)
            return process_fitz_doc(doc, dpi, quality)
        except Exception:
            pass
        try:
            return process_pillow_image(file_bytes, dpi, quality)
        except Exception:
            pass
        try:
            pdf_bytes = convert_doc_via_word(file_bytes, ext)
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            return process_fitz_doc(doc, dpi, quality)
        except Exception as e:
            raise RuntimeError(f"AI/EPS ফাইল প্রসেস করতে সমস্যা হয়েছে: {str(e)}")

    if ext in ['.doc', '.docx', '.rtf', '.odt', '.txt']:
        pdf_bytes = None

        # ১. দ্রুত Native Python Conversion প্রথমে চেষ্টা (Docx & Txt এর জন্য)
        if ext == '.docx':
            pdf_bytes = fallback_doc_to_pdf(file_bytes, ext)
        elif ext == '.txt':
            pdf_bytes = fallback_txt_to_pdf(file_bytes)

        # ২. Native না থাকলে বা ব্যর্থ হলে MS Word COM
        if not pdf_bytes:
            try:
                pdf_bytes = convert_doc_via_word(file_bytes, ext)
            except Exception as e:
                print(f"Word COM error for {ext}: {e}")

        # ৩. COM ব্যর্থ হলে চূড়ান্ত ফলব্যাক
        if not pdf_bytes and ext in ['.docx', '.doc', '.rtf', '.odt']:
            pdf_bytes = fallback_doc_to_pdf(file_bytes, ext)

        if pdf_bytes:
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            return process_fitz_doc(doc, dpi, quality)
        else:
            raise RuntimeError(f"{ext} ডকুমেন্টস ফাইল কনভার্ট করা সম্ভব হয়নি।")

    if ext in ['.xls', '.xlsx', '.csv', '.xlsm', '.ods']:
        pdf_bytes = None

        # ১. দ্রুত Native pandas/reportlab Conversion প্রথমে চেষ্টা (XLSX, CSV, ODS)
        if ext in ['.xlsx', '.csv', '.ods']:
            pdf_bytes = fallback_sheet_to_pdf(file_bytes, ext)

        # ২. Native না থাকলে বা ব্যর্থ হলে MS Excel COM
        if not pdf_bytes:
            try:
                pdf_bytes = convert_sheet_via_excel(file_bytes, ext)
            except Exception as e:
                print(f"Excel COM error for {ext}: {e}")

        # ৩. COM ব্যর্থ হলে ফলব্যাক
        if not pdf_bytes:
            pdf_bytes = fallback_sheet_to_pdf(file_bytes, ext)

        if pdf_bytes:
            doc = fitz.open(stream=pdf_bytes, filetype="pdf")
            return process_fitz_doc(doc, dpi, quality)
        else:
            raise RuntimeError(f"{ext} স্প্রেডশীট ফাইল কনভার্ট করা সম্ভব হয়নি।")

    raise ValueError(f"অসমর্থিত ফাইল ফরম্যাট: {ext}")

# --- ENGINE 2: CONVERT TO PDF ---

def convert_any_file_to_pdf(file_bytes, filename):
    _, ext = os.path.splitext(filename)
    ext = ext.lower()
    if not ext:
        ext = '.pdf'

    # 1. PDF
    if ext == '.pdf':
        return file_bytes

    # 2. SVG
    if ext == '.svg':
        doc = fitz.open(stream=file_bytes, filetype="svg")
        pdf_bytes = doc.convert_to_pdf()
        doc.close()
        return pdf_bytes

    # 3. Image formats (.jpeg, .jpg, .png, .webp, .gif, .psd)
    if ext in ['.jpeg', '.jpg', '.png', '.webp', '.gif', '.psd']:
        try:
            return process_pillow_to_pdf(file_bytes)
        except Exception:
            doc = fitz.open(stream=file_bytes)
            pdf_bytes = doc.convert_to_pdf()
            doc.close()
            return pdf_bytes

    # 4. Vector Graphics (.ai, .eps)
    if ext in ['.ai', '.eps']:
        try:
            doc = fitz.open(stream=file_bytes)
            pdf_bytes = doc.convert_to_pdf()
            doc.close()
            return pdf_bytes
        except Exception:
            pass
        try:
            return process_pillow_to_pdf(file_bytes)
        except Exception:
            pass
        try:
            return convert_doc_via_word(file_bytes, ext)
        except Exception as e:
            raise RuntimeError(f"AI/EPS ফাইল থেকে পিডিএফ রূপান্তর সম্ভব হয়নি: {e}")

    # 5. Documents (.doc, .docx, .rtf, .odt, .txt)
    if ext in ['.doc', '.docx', '.rtf', '.odt', '.txt']:
        pdf_bytes = None

        if ext == '.docx':
            pdf_bytes = fallback_doc_to_pdf(file_bytes, ext)
        elif ext == '.txt':
            pdf_bytes = fallback_txt_to_pdf(file_bytes)

        if not pdf_bytes:
            try:
                pdf_bytes = convert_doc_via_word(file_bytes, ext)
            except Exception as e:
                print(f"Word COM error for PDF conversion ({ext}): {e}")

        if not pdf_bytes and ext in ['.docx', '.doc', '.rtf', '.odt']:
            pdf_bytes = fallback_doc_to_pdf(file_bytes, ext)

        if pdf_bytes:
            return pdf_bytes
        else:
            raise RuntimeError(f"{ext} ডকুমেন্টস ফাইল থেকে পিডিএফ তৈরি সম্ভব হয়নি।")

    # 6. Spreadsheets (.xls, .xlsx, .csv, .xlsm, .ods)
    if ext in ['.xls', '.xlsx', '.csv', '.xlsm', '.ods']:
        pdf_bytes = None

        if ext in ['.xlsx', '.csv', '.ods']:
            pdf_bytes = fallback_sheet_to_pdf(file_bytes, ext)

        if not pdf_bytes:
            try:
                pdf_bytes = convert_sheet_via_excel(file_bytes, ext)
            except Exception as e:
                print(f"Excel COM error for PDF conversion ({ext}): {e}")

        if not pdf_bytes:
            pdf_bytes = fallback_sheet_to_pdf(file_bytes, ext)

        if pdf_bytes:
            return pdf_bytes
        else:
            raise RuntimeError(f"{ext} স্প্রেডশীট ফাইল থেকে পিডিএফ তৈরি সম্ভব হয়নি।")

    raise ValueError(f"অসমর্থিত ফাইল ফরম্যাট: {ext}")

# --- ROUTES FOR CONVERT TO JPG ---

@app.route('/convert_pdf', methods=['POST'])
@app.route('/convert_file', methods=['POST'])
def convert_pdf():
    files = request.files.getlist('pdf_file')
    if not files or files[0].filename == '':
        files = request.files.getlist('file')
        
    if not files or not any(f.filename for f in files):
        return jsonify({'error': 'দয়া করে অন্তত একটি ফাইল আপলোড করুন'}), 400
        
    dpi = int(request.form.get('dpi', 150))
    dpi = max(100, min(dpi, 600))
    
    file_tuples = []
    invalid_files = []
    
    for f in files:
        if not f or not f.filename:
            continue
        _, ext = os.path.splitext(f.filename)
        ext = ext.lower()
        if ext not in SUPPORTED_EXTENSIONS:
            invalid_files.append(f.filename)
        else:
            file_tuples.append((f.filename, f.read()))
            
    if invalid_files and not file_tuples:
        return jsonify({'error': f'অসমর্থিত ফাইল ফরম্যাট: {", ".join(invalid_files)}'}), 400
        
    if not file_tuples:
        return jsonify({'error': 'কোনো বৈধ ফাইল পাওয়া যায়নি'}), 400

    try:
        all_pages_data = []
        zip_buffer = io.BytesIO()
        global_page_counter = 1

        def process_file_item(args):
            file_idx, filename, file_bytes = args
            base_name = os.path.splitext(filename)[0]
            base_name_clean = WHITESPACE_PATTERN.sub('_', base_name)
            base_name_clean = CLEAN_FILENAME_PATTERN.sub('', base_name_clean).strip() or f"file_{file_idx+1}"
            pages_data, _ = convert_any_file_to_jpg(file_bytes, filename, dpi=dpi, quality=85)
            return file_idx, filename, base_name, base_name_clean, pages_data

        if len(file_tuples) > 1:
            with ThreadPoolExecutor(max_workers=min(MAX_WORKERS, len(file_tuples))) as executor:
                file_results = list(executor.map(process_file_item, [(i, name, b) for i, (name, b) in enumerate(file_tuples)]))
        else:
            file_results = [process_file_item((0, file_tuples[0][0], file_tuples[0][1]))]

        file_results.sort(key=lambda x: x[0])
        
        with zipfile.ZipFile(zip_buffer, 'w', zipfile.ZIP_STORED) as zip_file:
            for file_idx, filename, base_name, base_name_clean, pages_data in file_results:
                for p in pages_data:
                    if len(file_tuples) > 1:
                        img_filename = f"{base_name_clean}_page_{p['page']}.jpg"
                        label = f"{base_name} - পেজ {p['page']}"
                    else:
                        img_filename = f"page_{p['page']}.jpg"
                        label = f"পেজ/ছবি {p['page']}"
                    
                    b64_data = p['data_url'].split(',', 1)[1]
                    img_bytes = base64.b64decode(b64_data)
                    
                    zip_file.writestr(img_filename, img_bytes)
                    
                    all_pages_data.append({
                        'page': global_page_counter,
                        'file_page': p['page'],
                        'source_file': filename,
                        'filename': img_filename,
                        'label': label,
                        'data_url': p['data_url'],
                        'width': p['width'],
                        'height': p['height']
                    })
                    global_page_counter += 1
                    
        zip_buffer.seek(0)
        zip_b64 = base64.b64encode(zip_buffer.getvalue()).decode('utf-8')
        
        if len(file_tuples) == 1:
            first_name = os.path.splitext(file_tuples[0][0])[0]
            summary_name = WHITESPACE_PATTERN.sub(' ', first_name)
            summary_name = CLEAN_FILENAME_PATTERN.sub('', summary_name).strip() or "converted_file"
            ext_display = os.path.splitext(file_tuples[0][0])[1].upper().lstrip('.')
        else:
            summary_name = f"Combined_{len(file_tuples)}_Files"
            ext_display = "MULTI"
            
        return jsonify({
            'success': True,
            'pdf_name': summary_name,
            'total_files': len(file_tuples),
            'ext': ext_display,
            'total_pages': len(all_pages_data),
            'dpi': dpi,
            'pages': all_pages_data,
            'zip_b64': zip_b64
        })
    except Exception as e:
        print(f"Error converting files to JPG: {e}")
        return jsonify({'error': f"ফাইল কনভার্ট করতে সমস্যা হয়েছে: {str(e)}"}), 500

@app.route('/download_pdf_zip', methods=['POST'])
def download_pdf_zip():
    try:
        data = request.json or {}
        zip_b64 = data.get('zip_b64', '')
        pdf_name = data.get('pdf_name', 'converted_images')
        pages = data.get('pages', [])

        output = io.BytesIO()

        if pages:
            # Drag & Drop পেজ অর্ডার অনুসারে নতুন জিপ ফাইল তৈরি
            with zipfile.ZipFile(output, 'w', zipfile.ZIP_STORED) as zip_file:
                for idx, page in enumerate(pages):
                    data_url = page.get('data_url', '')
                    if ',' in data_url:
                        b64_str = data_url.split(',', 1)[1]
                        img_bytes = base64.b64decode(b64_str)
                        filename = f"page_{idx + 1}.jpg"
                        zip_file.writestr(filename, img_bytes)
        elif zip_b64:
            zip_bytes = base64.b64decode(zip_b64)
            output.write(zip_bytes)
        else:
            return jsonify({'error': 'কোনো ডাউনলোডেবল ডাটা পাওয়া যায়নি'}), 400

        output.seek(0)
        pdf_name_clean = WHITESPACE_PATTERN.sub(' ', pdf_name)
        pdf_name_clean = CLEAN_FILENAME_PATTERN.sub('', pdf_name_clean).strip() or "converted_images"
        
        return send_file(
            output,
            mimetype='application/zip',
            as_attachment=True,
            download_name=f"{pdf_name_clean}_images.zip"
        )
    except Exception as e:
        return jsonify({'error': f"জিপ ডাউনলোড ত্রুটি: {str(e)}"}), 500

# --- ROUTES FOR CONVERT TO PDF ---

@app.route('/convert_to_pdf', methods=['POST'])
def convert_to_pdf_route():
    files = request.files.getlist('file')
    if not files or files[0].filename == '':
        files = request.files.getlist('pdf_file')
        
    if not files or not any(f.filename for f in files):
        return jsonify({'error': 'দয়া করে অন্তত একটি ফাইল আপলোড করুন'}), 400
        
    file_tuples = []
    invalid_files = []
    
    for f in files:
        if not f or not f.filename:
            continue
        _, ext = os.path.splitext(f.filename)
        ext = ext.lower()
        if not ext:
            ext = '.pdf'
        if ext not in SUPPORTED_EXTENSIONS:
            invalid_files.append(f.filename)
        else:
            file_tuples.append((f.filename, f.read()))
            
    if invalid_files and not file_tuples:
        return jsonify({'error': f'অসমর্থিত ফাইল ফরম্যাট: {", ".join(invalid_files)}'}), 400
        
    if not file_tuples:
        return jsonify({'error': 'কোনো বৈধ ফাইল পাওয়া যায়নি'}), 400
        
    try:
        def process_single_pdf(args):
            file_idx, filename, file_bytes = args
            sub_pdf_bytes = convert_any_file_to_pdf(file_bytes, filename)
            return file_idx, sub_pdf_bytes

        if len(file_tuples) > 1:
            with ThreadPoolExecutor(max_workers=min(MAX_WORKERS, len(file_tuples))) as executor:
                pdf_results = list(executor.map(process_single_pdf, [(i, name, b) for i, (name, b) in enumerate(file_tuples)]))
        else:
            pdf_results = [process_single_pdf((0, file_tuples[0][0], file_tuples[0][1]))]

        pdf_results.sort(key=lambda x: x[0])

        doc_master = fitz.open()
        for _, sub_pdf_bytes in pdf_results:
            sub_doc = fitz.open(stream=sub_pdf_bytes, filetype="pdf")
            doc_master.insert_pdf(sub_doc)
            sub_doc.close()
            
        combined_pdf_bytes = doc_master.tobytes(clean=True, deflate=True)
        
        # Parallel preview generation at 120 DPI
        def render_preview_page(i):
            page = doc_master.load_page(i)
            pix = page.get_pixmap(dpi=120, alpha=False)
            img_bytes = pix.tobytes("jpeg", jpg_quality=80)
            b64_str = base64.b64encode(img_bytes).decode('utf-8')
            return i, {
                'page': i + 1,
                'data_url': f"data:image/jpeg;base64,{b64_str}",
                'width': pix.width,
                'height': pix.height
            }

        total_pages = len(doc_master)
        if total_pages > 1:
            with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
                preview_results = list(executor.map(render_preview_page, range(total_pages)))
        else:
            preview_results = [render_preview_page(0)]

        preview_results.sort(key=lambda x: x[0])
        preview_pages = [p for _, p in preview_results]

        doc_master.close()
        
        pdf_b64 = base64.b64encode(combined_pdf_bytes).decode('utf-8')
        
        if len(file_tuples) == 1:
            first_name = os.path.splitext(file_tuples[0][0])[0]
            summary_name = WHITESPACE_PATTERN.sub(' ', first_name)
            summary_name = CLEAN_FILENAME_PATTERN.sub('', summary_name).strip() or "converted_document"
            ext_display = os.path.splitext(file_tuples[0][0])[1].upper().lstrip('.')
        else:
            summary_name = f"Combined_{len(file_tuples)}_Files"
            ext_display = "MULTI"
            
        return jsonify({
            'success': True,
            'file_name': summary_name,
            'total_files': len(file_tuples),
            'ext': ext_display,
            'total_pages': total_pages,
            'pdf_size': len(combined_pdf_bytes),
            'pdf_b64': pdf_b64,
            'preview_pages': preview_pages
        })
    except Exception as e:
        print(f"Error converting to PDF: {e}")
        return jsonify({'error': f"পিডিএফ তৈরি করতে সমস্যা হয়েছে: {str(e)}"}), 500

@app.route('/download_pdf_file', methods=['POST'])
def download_pdf_file():
    try:
        data = request.json or {}
        pdf_b64 = data.get('pdf_b64', '')
        file_name = data.get('file_name', 'converted_document')
        pages = data.get('pages', [])

        if pages:
            # Drag & Drop পেজ অর্ডার অনুসারে নতুন পিডিএফ ফাইল তৈরি
            doc_master = fitz.open()
            for page in pages:
                data_url = page.get('data_url', '')
                if ',' in data_url:
                    b64_str = data_url.split(',', 1)[1]
                    img_bytes = base64.b64decode(b64_str)
                    
                    img_doc = fitz.open(stream=img_bytes)
                    pdf_bytes = img_doc.convert_to_pdf()
                    img_doc.close()
                    
                    sub_doc = fitz.open(stream=pdf_bytes, filetype="pdf")
                    doc_master.insert_pdf(sub_doc)
                    sub_doc.close()
                    
            combined_pdf_bytes = doc_master.tobytes(clean=True, deflate=True)
            doc_master.close()
            output = io.BytesIO(combined_pdf_bytes)
        elif pdf_b64:
            pdf_bytes = base64.b64decode(pdf_b64)
            output = io.BytesIO(pdf_bytes)
        else:
            return jsonify({'error': 'কোনো ডাউনলোডেবল ডাটা পাওয়া যায়নি'}), 400

        output.seek(0)
        file_name_clean = WHITESPACE_PATTERN.sub(' ', file_name)
        file_name_clean = CLEAN_FILENAME_PATTERN.sub('', file_name_clean).strip() or "converted_document"
        
        return send_file(
            output,
            mimetype='application/pdf',
            as_attachment=True,
            download_name=f"{file_name_clean}.pdf"
        )
    except Exception as e:
        return jsonify({'error': f"পিডিএফ ডাউনলোড ত্রুটি: {str(e)}"}), 500

if __name__ == '__main__':
    app.run(host='0.0.0.0', port=5000, debug=True)